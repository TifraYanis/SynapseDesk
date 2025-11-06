"""
generate_rich_testdata.py
Génère un jeu de données synthétiques professionnels variés pour test RAG.
Crée des fichiers dans api/data/ : .md, .txt, .docx, .json, .csv et test_queries.jsonl
"""

import os, json, random, datetime, textwrap
from pathlib import Path

try:
    from docx import Document
except Exception:
    Document = None  # si python-docx non installé, on écrira seulement txt/md/json

BASE = Path(__file__).resolve().parent
DATA_DIR = BASE / "data"
DATA_DIR.mkdir(exist_ok=True)

now = datetime.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"

# --- roles & personas ---
personas = [
    {"name": "Alice Martin", "role": "Data Scientist", "email": "alice.martin@example.com"},
    {"name": "Bastien Leroy", "role": "Site Reliability Engineer", "email": "bastien.leroy@example.com"},
    {"name": "Camille Dupont", "role": "Product Owner", "email": "camille.dupont@example.com"},
    {"name": "David Moreau", "role": "Data Engineer", "email": "david.moreau@example.com"},
    {"name": "Elisa Bernard", "role": "HR Manager", "email": "elisa.bernard@example.com"},
    {"name": "Fouad K.", "role": "Support Technician", "email": "fouad.k@example.com"},
    {"name": "Gwen Thomas", "role": "Business Analyst", "email": "gwen.thomas@example.com"},
    {"name": "Hugo Petit", "role": "ML Ops Engineer", "email": "hugo.petit@example.com"}
]

# --- helper writers ---
def write_text(path: Path, text: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")
    print("Wrote", path)

def write_docx(path: Path, title: str, paragraphs: list):
    if Document is None:
        # fallback to txt
        content = title + "\n\n" + "\n\n".join(paragraphs)
        write_text(path.with_suffix(".txt"), content)
        return
    doc = Document()
    doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print("Wrote", path)

# --- generate markdown knowledge base ---
kb_md = {
    "Troubleshooting_Spark.md": textwrap.dedent("""\
        # Troubleshooting Spark

        **Contexte**: jobs ETL nocturnes, cluster YARN, version Spark 3.4.

        ## Problèmes communs
        - OOM (OutOfMemory) sur executors: augmenter `spark.executor.memory` ou réduire `spark.executor.cores`.
        - Shuffle spill: vérifier partitioning, diminuer `spark.sql.shuffle.partitions`.
        - Skew: appliquer salting sur les clefs chaudes, ou rééquilibrer données.

        ## Actions recommandées
        1. Rejouer le job sur un dataset réduit.
        2. Activer `spark.sql.adaptive.enabled = true`.
        3. Ajouter monitoring: métriques GC, memory, shuffle write.

        **Contacts**: David Moreau (Data Engineer), bastien.leader@example.com
        """),
    "DataGovernance_Policy.md": textwrap.dedent("""\
        # Politique de Gouvernance des Données

        - Classification: PUBLIC / INTERNAL / SENSITIVE
        - PII handling: anonymisation obligatoire avant partage externe.
        - Retention: logs opérationnels 90 jours, backups 1 an.

        **Processus**: requests to data access via Jira + validation RH/Legal.
        """),
    "PO_Roadmap.md": textwrap.dedent("""\
        # Roadmap Produit (Q4)

        ## Épics
        - Observabilité pipeline (priority: high)
        - Interface usage analytics (priority: medium)
        - Automatisation des déploiements (priority: low)

        ## Notes
        PO: Camille Dupont. Sprint planning chaque lundi 09:30.
        """)
}

for name, text in kb_md.items():
    write_text(DATA_DIR / name, text)

# --- generate docx incident reports & runbooks ---
incident_paragraphs = [
    "Incident: job bronze_to_silver failed with ExecutorLostFailure (OOM).",
    "Time: 2024-09-12 02:33 UTC. Cluster: prod-east-1.",
    "Diagnostic: spike de données sur partition date=2024-09-11, skew détecté.",
    "Action: redémarrer job sur small-sample, augmenter shuffle partitions, contacter oncall."
]
write_docx(DATA_DIR / "Incident_OOM.docx", "Incident Report — OOM", incident_paragraphs)

runbook_paragraphs = [
    "Runbook: Quick recovery steps for failed ETL jobs",
    "1) Check Spark UI for failed stages",
    "2) If OOM, adjust executor memory and re-run with smaller parallelism",
    "3) Notify data-platform Slack channel #dataops"
]
write_docx(DATA_DIR / "Runbook_ETL.docx", "Runbook ETL — Recovery", runbook_paragraphs)

# --- generate HR policy + onboarding text files ---
hr_text = textwrap.dedent(f"""\
    ONBOARDING CHECKLIST
    - Name: New hire
    - Team: Data Platform
    - Mentor: {personas[4]['name']} ({personas[4]['email']})
    - Tools access: GitHub, Jenkins, Confluence, internal SSO
    ---
    Confidential: Salary details are handled by payroll and not stored in Confluence.
    """)
write_text(DATA_DIR / "Onboarding_Checklist.txt", hr_text)

hr_policy = textwrap.dedent("""\
    HR Policy - Remote Work
    - Remote days allowed: up to 3 days/week
    - Equipment: laptop provided, monitor allowance
    - Contact HR for ergonomic requests
    """)
write_text(DATA_DIR / "HR_RemotePolicy.md", hr_policy)

# --- generate product spec (PO) ---
prod_spec = textwrap.dedent(f"""\
    # Feature Spec: Usage Analytics Dashboard

    Author: {personas[2]['name']} ({personas[2]['email']})
    Date: {now}

    ## Goal
    Provide product managers with daily active usage metrics for the data portal.

    ## Requirements
    - Daily ingestion of web logs
    - Aggregation by user cohort
    - Export as CSV and alert on drops > 20%

    ## Acceptance criteria
    - Dashboard available within 2 weeks
    """)
write_text(DATA_DIR / "Feature_UsageAnalytics.md", prod_spec)

# --- generate email-like notes ---
emails = [
    {
        "subject": "Action required: nightly job failures",
        "from": "ops@example.com",
        "to": "data-platform@example.com",
        "date": now,
        "body": "We've observed increased failures for bronze_to_silver during peak. Please investigate and escalate if needed."
    },
    {
        "subject": "New hire: access request",
        "from": personas[4]['email'],
        "to": "it-ops@example.com",
        "date": now,
        "body": "Please create accounts: GitHub, Jira, Confluence for new hire - onboarding next Monday."
    }
]
for i, m in enumerate(emails, 1):
    write_text(DATA_DIR / f"email_{i:02d}.json", json.dumps(m, ensure_ascii=False, indent=2))

# --- generate technical code snippets and configs ---
code_samples = {
    "bronze_to_silver_job.py": textwrap.dedent("""\
        from pyspark.sql import SparkSession
        spark = SparkSession.builder.appName("BronzeToSilver").getOrCreate()
        df = spark.read.format("delta").load("/data/bronze")
        df = df.filter("status = 'active'")
        df.write.mode("overwrite").format("delta").save("/data/silver")
        """),
    "spark_conf.properties": textwrap.dedent("""\
        spark.executor.memory=8g
        spark.executor.cores=2
        spark.sql.shuffle.partitions=200
        """)
}
for name, content in code_samples.items():
    write_text(DATA_DIR / name, content)

# --- generate logs (JSON) ---
job_ids = ["bronze_to_silver","client_merge","daily_agg"]
errors = [None, "ExecutorLostFailure (OOM)", "StageFailure (Skew detected)", "FileNotFoundError"]
for i in range(1,201):
    rec = {
        "log_id": i,
        "job_id": random.choice(job_ids),
        "stage": random.randint(1,8),
        "duration_ms": random.randint(200, 300000),
        "error": random.choice(errors),
        "host": f"node-{random.randint(1,20)}",
        "timestamp": (datetime.datetime.utcnow() - datetime.timedelta(minutes=random.randint(0,10000))).isoformat() + "Z"
    }
    write_text(DATA_DIR / f"logs/log_{i:04d}.json", json.dumps(rec, ensure_ascii=False))

# --- generate jira-like tickets ---
tickets = []
for i in range(1,61):
    t = {
        "ticket_id": f"DP-{1000+i}",
        "summary": random.choice([
            "Job bronze_to_silver fails intermittently",
            "Add alert for low disk on prod cluster",
            "Data schema version bump causes downstream errors",
            "Request access to analytics dashboard"
        ]),
        "reporter": random.choice(personas)["email"],
        "assignee": random.choice(personas)["email"],
        "priority": random.choice(["Low","Medium","High","Critical"]),
        "status": random.choice(["Open","In Progress","Resolved","Closed"]),
        "created_at": (datetime.datetime.utcnow() - datetime.timedelta(days=random.randint(0,60))).isoformat() + "Z",
        "description": "Steps to reproduce: ..."
    }
    tickets.append(t)
    write_text(DATA_DIR / f"tickets/{t['ticket_id']}.json", json.dumps(t, ensure_ascii=False, indent=2))

# --- generate small CSV inventory (datasets) ---
csv_lines = ["dataset_name,owner,rows,updated_at"]
for name in ["customers","transactions","events","products","inventory"]:
    csv_lines.append(f"{name},{random.choice(personas)['email']},{random.randint(1000,1000000)},{now}")
write_text(DATA_DIR / "datasets_catalog.csv", "\n".join(csv_lines))

# --- generate FAQ and how-to guides ---
faq = textwrap.dedent("""\
    Q: Who owns bronze_to_silver job?
    A: data-eng@example.com (David Moreau)

    Q: How to contact oncall?
    A: Pager duty -> pager@example.com
    """)
write_text(DATA_DIR / "FAQ.txt", faq)

howto = textwrap.dedent("""\
    How to add a new dataset to the pipeline:
    1) Create schema PR in repo
    2) Update ingestion config in /configs/ingestion.yaml
    3) Run local tests and push
    """)
write_text(DATA_DIR / "HowTo_AddDataset.md", howto)

# --- generate personas file (useful for testing) ---
write_text(DATA_DIR / "personas.json", json.dumps(personas, ensure_ascii=False, indent=2))

# --- generate test queries (contextual vs non-contextual) ---
test_queries = [
    {"id":"q001","query":"Comment corriger un OOM sur un job Spark ?","type":"contextual","expected":"référence aux actions: augmenter executor memory, partitioning, broadcast join"},
    {"id":"q002","query":"Qui est le PO pour la roadmap usage analytics ?","type":"contextual","expected":"Camille Dupont"},
    {"id":"q003","query":"Quelles sont les règles pour le télétravail ?","type":"contextual","expected":"Remote days allowed: up to 3 days/week"},
    {"id":"q004","query":"Comment reset un mot de passe AD ?","type":"non_contextual","expected":"réponse générique procédure IT (pas dans le KB)"},
    {"id":"q005","query":"Quel est le propriétaire du dataset transactions ?","type":"contextual","expected":"email du owner dans datasets_catalog.csv"},
    {"id":"q006","query":"Explique-moi le concept de 'data skew' en Spark simplement.","type":"non_contextual","expected":"explication pédagogique sans citation spécifique"},
    {"id":"q007","query":"Le job bronze_to_silver a échoué la nuit dernière, quelles actions ?","type":"contextual","expected":"Analyse logs + OOM + suggestion: rerun with small sample"},
    {"id":"q008","query":"Comment contacter l'oncall ?","type":"contextual","expected":"pager@example.com ou data-platform@example.com"}
]
with open(DATA_DIR / "test_queries.jsonl", "w", encoding="utf-8") as f:
    for q in test_queries:
        f.write(json.dumps(q, ensure_ascii=False) + "\n")
print("Wrote test_queries.jsonl")

print("Synthetic professional dataset generated in:", DATA_DIR)
